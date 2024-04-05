from PyQt5 import QtWidgets, QtGui, QtCore
from untitled1 import Ui_MainWindow
from add_worker import Ui_MainWindow as add_workerses
from more_workers import Ui_MainWindow as UI_more_workers
from add_offer import Ui_MainWindow as Ui_add_offer
from more_offer import Ui_MainWindow as Ui_more_offer
from add_service import Ui_MainWindow as Ui_add_service
from add_sale import Ui_MainWindow as Ui_add_sale
from more_service import Ui_MainWindow as Ui_more_service
from more_sale import Ui_MainWindow as Ui_more_sale
from new_service import Ui_MainWindow as Ui_new_service
from new_sale import Ui_MainWindow as Ui_new_sale
from add_rash import Ui_MainWindow as Ui_add_rash
from add_group import Ui_MainWindow as Ui_add_group
from more_rash import Ui_MainWindow as Ui_more_rash
from stat_change import Ui_MainWindow as Ui_stat_change
from first_start import Ui_MainWindow as Ui_first_start
from rec import Ui_MainWindow as Ui_rec
from db_working import *
from datetime import datetime
import datetime as dt
import sys
from docx import Document
from docx.shared import Pt
import os
import getpass

windows = []
passwad = "H7)%i5&ucwQs#a02"
latest = dt.date(3228, 3, 10)
today = dt.datetime.now().date()
passww = load_pass()




def autho_check():
    f = open(f"C:/Users/{getpass.getuser()}/AppData/Local/myprogrampy/oleg", 'r')
    if f.read() == passwad:
        f.close()
        return True
    else:
        f.close()
        return False


class Add_group(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        super(Add_group, self).__init__()
        self.ui = Ui_add_group()
        self.ui.setupUi(self)
        self.ui.bt_add.clicked.connect(self.confirm)
        self.ui.bt_close.clicked.connect(self.close)

    def confirm(self):
        if self.ui.line_sur.text().strip() != '':
            add_group(self.ui.line_sur.text().strip())
            self.g.services_open()
            self.g.ui.service_group.addItem(self.ui.line_sur.text().strip())
            self.close()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)
        print(windows)


class New_service(QtWidgets.QMainWindow):
    def __init__(self, g, f, offer_number):
        self.f = f
        self.g = g
        self.offer_number = offer_number
        ssss = load_offers()
        type = 'l'
        ss = load_services1([offer_number])
        services_in_offer = {}
        services_sost = {}
        self.services_comm = {}
        for i in ss:
            services_in_offer[i[1]] = i[3]
            services_sost[i[1]] = i[6]
            self.services_comm[i[1]] = i[7]
        for i in ssss:
            if int(i[0]) == int(offer_number):
                type = i[1]
        self.type = type
        super(New_service, self).__init__()
        self.ui = Ui_new_service()
        self.ui.setupUi(self)
        date = str(datetime.now().date())
        self.ui.l_date.setText(self.ui.l_date.text() + date)
        self.services_in_offer = services_in_offer
        self.services_sost = services_sost
        self.ui.groups.addItem('Все')
        k = load_groups()
        for i in k:
            self.ui.groups.addItem(i[0])
        self.ui.groups.currentTextChanged.connect(self.update_table)
        self.ui.bt_confirm.clicked.connect(self.close)
        self.ui.bt_cancel.clicked.connect(self.close)
        self.update_table()


    def update_table(self):
        print(self.services_sost)
        r = self.ui.tableWidget.rowCount()
        for i in range(r):
            self.ui.tableWidget.removeRow(0)
        services_in_offer = self.services_in_offer
        if self.ui.groups.currentText() == 'Все':
            services = load_services()
        else:
            services = load_services(self.ui.groups.currentText())
        services_sost = self.services_sost
        services = list(sorted(services, key=lambda s: s[1]))
        self.ui.tableWidget.setRowCount(len(services))
        for i in range(len(services)):
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(services[i][1]))
            ol = QtGui.QColor(255, 255, 255)
            zxc = 2
            if self.type == 'g':
                zxc = 4
            col = 1
            ff = ''
            if services[i][1] in self.services_in_offer:
                ff = str(self.services_comm[services[i][1]])
                if services_sost[services[i][1]]:
                    ol = QtGui.QColor(100, 100, 100)
                else:
                    ol = QtGui.QColor(100, 255, 100)
                col = services_in_offer[services[i][1]]
            self.ui.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(services[i][zxc]))
            self.ui.tableWidget.setItem(i, 2, QtWidgets.QTableWidgetItem(str(col)))
            self.ui.tableWidget.item(i, 2).setBackground(ol)
            self.ui.tableWidget.setItem(i, 3, QtWidgets.QTableWidgetItem(ff))
            bt = QtWidgets.QPushButton('Добавить')
            bt.offer_number = i
            bt.name_of = services[i][1]
            bt.clicked.connect(self.confirm)
            self.ui.tableWidget.setCellWidget(i, 4, bt)
        self.ui.tableWidget.resizeColumnsToContents()

    def confirm(self):
        i = QtWidgets.QApplication.instance().sender()
        z = i.name_of
        k = 1
        if z in self.services_sost:
            if self.services_sost[z] == 1:
                k = 0
        if k:
            i = i.offer_number
            date = str(datetime.now().date())
            if self.ui.tableWidget.item(i, 2).text().strip() == '':
                self.ui.tableWidget.item(i, 2).setText('0')
            if self.ui.tableWidget.item(i, 1).text().strip() == '':
                self.ui.tableWidget.item(i, 1).setText('0')
            update_offer(self.offer_number, work=z, price=str(float(self.ui.tableWidget.item(i, 1).text().replace(',', '.')) * int(self.ui.tableWidget.item(i, 2).text())).replace('.', ','), dates=date, worker='', col=self.ui.tableWidget.item(i, 2).text(), comm=self.ui.tableWidget.item(i, 3).text())
            self.ui.tableWidget.item(i, 2).setBackground(QtGui.QColor(100, 255, 100))
            self.services_in_offer[z] = self.ui.tableWidget.item(i, 2).text()
            self.services_sost[z] = 0
            self.services_comm[z] = self.ui.tableWidget.item(i, 3).text()
            self.f.update_table()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)
        self.g.offers_open()


class New_sale(QtWidgets.QMainWindow):
    def __init__(self, g, f, offer_number):
        self.f = f
        self.g = g
        self.offer_number = offer_number
        ssss = load_offers()
        type = 'l'
        ss = load_sales([offer_number])
        services_in_offer = {}
        for i in ss:
            services_in_offer[i[1]] = i[3]
        self.type = type
        super(New_sale, self).__init__()
        self.ui = Ui_new_sale()
        self.ui.setupUi(self)
        date = str(datetime.now().date())
        self.ui.l_date.setText(self.ui.l_date.text() + date)
        services = load_in()
        services = list(sorted(services, key=lambda s: s[1]))
        self.ui.tableWidget.setRowCount(len(services))
        for i in range(len(services)):
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(services[i][1]))
            ol = QtGui.QColor(255, 255, 255)
            col = 1
            if services[i][1] in services_in_offer:
                col = services_in_offer[services[i][1]]
                ol = QtGui.QColor(100, 255, 100)
            self.ui.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(str(services[i][3])))
            self.ui.tableWidget.setItem(i, 2, QtWidgets.QTableWidgetItem(str(col)))
            self.ui.tableWidget.item(i, 2).setBackground(ol)
            bt = QtWidgets.QPushButton('Добавить')
            bt.offer_number = i
            bt.name_of = services[i][1]
            bt.clicked.connect(self.confirm)
            self.ui.tableWidget.setCellWidget(i, 3, bt)
        self.ui.bt_confirm.clicked.connect(self.close)
        self.ui.bt_cancel.clicked.connect(self.close)
        self.ui.tableWidget.resizeColumnsToContents()

    def confirm(self):
        i = QtWidgets.QApplication.instance().sender()
        z = i.name_of
        i = i.offer_number
        date = str(datetime.now().date())
        if self.ui.tableWidget.item(i, 2).text().strip() == '':
            self.ui.tableWidget.item(i, 2).setText('0')
        if self.ui.tableWidget.item(i, 1).text().strip() == '':
            self.ui.tableWidget.item(i, 1).setText('0')
        update_offer_sale(self.offer_number, work=z, price=str(float(self.ui.tableWidget.item(i, 1).text().replace(',', '.')) * int(self.ui.tableWidget.item(i, 2).text())).replace('.', ','), dates=date, col=self.ui.tableWidget.item(i, 2).text())
        self.ui.tableWidget.item(i, 2).setBackground(QtGui.QColor(100, 255, 100))
        self.f.update_table()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)
        self.g.offers_open()


class First_start_w(QtWidgets.QMainWindow):
    def __init__(self):
        super(First_start_w, self).__init__()
        self.ui = Ui_first_start()
        self.ui.setupUi(self)
        self.ui.bt_add.clicked.connect(self.ok)
        self.ui.bt_close.clicked.connect(self.closeEvent)

    def ok(self):
        try:
            passw = self.ui.line_sur.text()
            passw = int(passw)
            save_pass(passw)
            self.close()
            application.show()
        except ValueError:
            self.ui.l_ex.setText('Присутствуют недопустимые значения')

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()


class Rec(QtWidgets.QMainWindow):
    def __init__(self, g, gg):
        self.gg = gg
        print(self.gg.number)
        self.g = g
        super(Rec, self).__init__()
        self.ui = Ui_rec()
        self.ui.setupUi(self)
        self.ui.btn_save_2.clicked.connect(self.close)
        self.ui.btn_save.clicked.connect(self.ok)
        offer = list(filter(lambda s: s[0] == g[0], load_offers()))[0]
        self.ui.t_rec.setText(offer[5])
        self.ui.t_os.setText(offer[4])

    def ok(self):
        update_rec(self.g[0], self.ui.t_rec.toPlainText(), self.ui.t_os.toPlainText())
        self.gg.number = list(self.gg.number)
        self.gg.number[5] = self.ui.t_rec.toPlainText()
        self.gg.number[4] = self.ui.t_os.toPlainText()
        self.close()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)


class Stat_change(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        self.passwo = ''
        self.btn = QtWidgets.QApplication.instance().sender()
        self.offer = self.btn.offer_number
        self.serv_info = self.btn.serv_info
        super(Stat_change, self).__init__()
        self.ui = Ui_stat_change()
        self.ui.setupUi(self)
        self.ui.btn_close.clicked.connect(self.close)
        self.ui.btn_com.clicked.connect(self.confirm)
        self.ui.l_password.hide()
        self.ui.label_password.hide()
        if self.serv_info[6]:
            self.ui.l_password.show()
            self.ui.label_password.show()
            self.ui.co_worker.addItem(self.serv_info[5])
            self.ui.co_worker.setEnabled(False)
            self.ui.btn_com.setText('Вернуть в работу')
        else:
            workers = load_workers()
            for i in workers:
                if i[3]:
                    self.ui.co_worker.addItem(' '.join(i[1].split('ALEEE')))
        self.ui.l_password.textEdited.connect(self.password_editing)

    def password_editing(self):
        len_1 = len(self.ui.l_password.text())
        len_2 = len(self.passwo)
        if len_1 < len_2:
            self.passwo = self.passwo[:(len_1-len_2)]
        elif len_1 > len_2:
            self.passwo = self.passwo + self.ui.l_password.text()[-1]
        self.ui.l_password.setText('*' * len_1)

    def confirm(self):
        if self.serv_info[6]:
            if self.passwo == str(load_pass()[0][0]):
                update_offer(self.serv_info[0], work=self.serv_info[1], price=self.serv_info[2], col=self.serv_info[3],
                             dates=self.serv_info[4], worker='', sost=0, comm=self.serv_info[7])
                delete_from_worker('ALEEE'.join(self.ui.co_worker.currentText().split(' ')), number=self.serv_info[0], works=self.serv_info[1])
                self.g.update_table()
                self.close()
        else:
            update_offer(self.serv_info[0], work=self.serv_info[1], price=self.serv_info[2],
                         col=self.serv_info[3], dates=self.serv_info[4],
                         worker=self.ui.co_worker.currentText(), sost=1, comm=self.serv_info[7])
            update_worker('ALEEE'.join(self.ui.co_worker.currentText().split(' ')), works=self.serv_info[1], prices=self.serv_info[2], dates=self.serv_info[4], number=self.serv_info[0])
            self.g.update_table()
            self.close()


    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)


class More_offer(QtWidgets.QMainWindow):
    def __init__(self, g, number):
        self.g = g
        self.passwo = ''
        self.number = number
        if number == []:
            self.btn = QtWidgets.QApplication.instance().sender()
            self.number = self.btn.offer[0]
        super(More_offer, self).__init__()
        self.ui = Ui_more_offer()
        self.ui.setupUi(self)
        type = self.number[1]
        if type == 'l':
            self.ui.l_type.addItem('Легковой')
            self.ui.l_type.addItem('Грузовой')
        else:
            self.ui.l_type.addItem('Грузовой')
            self.ui.l_type.addItem('Легковой')
        self.ui.label.setText(self.ui.label.text() + str(self.number[0]))
        self.ui.tableWidget.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self.ui.l_date_.setText(str(self.number[6]))
        self.ui.l_name_.setText(str(self.number[2]))
        self.ui.l_client_.setText(str(self.number[3]))
        self.ui.l_numberr.setText(str(self.number[8]))
        self.ui.btn_close_offer.hide()
        self.ui.l_password.setAutoFillBackground(True)
        self.ui.l_password.textEdited.connect(self.password_editing)
        self.ui.btn_add_service.clicked.connect(self.new_service)
        self.ui.btn_confirm.clicked.connect(self.confirm)
        self.ui.btn_sost.clicked.connect(self.update_sost)
        self.ui.btn_close_offer.clicked.connect(self.close_realy)
        self.ui.l_password.hide()
        self.ui.la_password.hide()
        self.ui.l_date_1.hide()
        self.ui.label_9.hide()
        self.ui.l_date_1.setText(self.number[9])
        self.closeee = True
        if self.number[7] == 'in':
            self.ui.pushButton.hide()
            self.ui.btn_sost.setText('В работе')
        elif self.number[7] == 'nw':
            self.ui.pushButton.hide()
            self.ui.btn_sost.setText('Новый')
        elif self.number[7] == 'ne':
            self.ui.pushButton.hide()
            self.ui.btn_sost.setText('Готов')
        elif self.number[7] == 'co' or self.number[7] == 'cor':
            self.ui.pushButton.hide()
            self.ui.btn_sost.setText('Закрыт')
            self.ui.btn_close_offer.show()
            self.ui.l_password.show()
            self.ui.la_password.show()
        elif self.number[7] == 'wa':
            self.ui.btn_sost.setText('В работе, ждёт запчастей')
            self.ui.pushButton.hide()
        elif self.number[7] == 'cl':
            self.ui.btn_sost.setText('Отменён')
            self.ui.la_password.show()
            self.ui.l_password.show()
            self.ui.pushButton.show()
        if self.number[7] == 'cor':
            self.ui.pushButton.hide()
            self.ui.l_date_1.show()
            self.ui.label_9.show()
            self.ui.btn_add_service.hide()
            self.ui.btn_sost.setEnabled(False)
            self.ui.btn_close_offer.hide()
            self.ui.l_password.hide()
            self.ui.la_password.hide()
            self.ui.btn_add_sale.hide()
            self.ui.l_client_.setEnabled(False)
            self.ui.l_type.setEnabled(False)
            self.ui.l_numberr.setEnabled(False)
            self.closeee = False
        self.sost_to_save = self.number[7]
        self.update_table()
        self.ui.btn_add_sale.clicked.connect(self.new_sale)
        self.ui.bt_print.clicked.connect(self.to_print)
        self.ui.btn_add_rec.clicked.connect(self.rec)
        self.ui.pushButton.clicked.connect(self.delete_off)

    def confirm(self):
        type = 'l'
        if self.ui.l_type.currentText()[0] == 'Г':
            type = 'g'
        update_offer_data(self.number[0], self.ui.l_client_.text(), self.ui.l_numberr.text(), type)
        update_sost(str(self.number[0]), self.sost_to_save)
        self.g.offers_open()
        self.close()

    def remove_row(self, table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    def to_print(self, number):
        number = list(self.number)
        print(number)
        if number[9] == None:
            doc = Document('offer_list.docx')
            table = doc.tables[0]
            cell = table.cell(0, 0)
            new_text = cell.text.split('\n')
            new_text[0] += str(number[0])
            new_text[1] += number[6]
            new_text[3] += number[3]
            new_text[4] += number[8]
            new_text[5] += number[2]
            cell.text = '\n'.join(new_text)
            table = doc.tables[1]
            row = table.rows[-1]
            ss = load_services1(number)
            for i in range(6):
                self.remove_row(table, table.rows[1])
            o = 0
            z = 0
            for i in range(len(ss)):
                if ss[i][3] != '0':
                    table.add_row()
                    k = table.cell(i + 2 - o, 0)
                    k.text = str(i + 1 - o)
                    k = table.cell(i + 2 - o, 1)
                    k.text = ss[i][1]
                    k = table.cell(i + 2 - o, 3)
                    k.text = ss[i][3]
                    k = table.cell(i + 2 - o, 4)
                    k.text = str(float(ss[i][2].replace(',', '.')) / int(ss[i][3])).replace('.', ',')
                    k = table.cell(i + 2 - o, 5)
                    k.text = ss[i][2]
                    z += float(ss[i][2].replace(',', '.'))
                else:
                    o += 1
            k = table.cell(1, 5)
            k.text = str(z).replace('.', ',')
            row0 = table.rows[1]
            row1 = table.rows[-1]
            row1._tr.addnext(row0._tr)
            p = doc.paragraphs[5]
            p.text = f'Расходная накладная к Заказу №{str(number[0])}'

            table = doc.tables[2]
            ss = load_sales(number)
            for i in range(4):
                self.remove_row(table, table.rows[1])
            z1 = z
            z = 0
            o = 0
            for i in range(len(ss)):
                if ss[i][3] != '0':
                    table.add_row()
                    k = table.cell(i + 2 - o, -1)
                    k.text = str(i + 1 - o)
                    k = table.cell(i + 2 - o, 0)
                    k.text = ss[i][1]
                    k = table.cell(i + 2 - o, 2)
                    k.text = ss[i][3]
                    k = table.cell(i + 2 - o, 3)
                    k.text = str(float(ss[i][2].replace(',', '.')) / int(ss[i][3])).replace('.', ',')
                    k = table.cell(i + 2 - o, 4)
                    k.text = ss[i][2]
                    k = table.cell(i + 2 - o, 5)
                    k.text = '0'
                    z += float(ss[i][2].replace(',', '.'))
                else:
                    o += 1
            k = table.cell(1, 5)
            k.text = str(z).replace('.', ',')
            z1 += z
            row0 = table.rows[1]
            row1 = table.rows[-1]
            row1._tr.addnext(row0._tr)
            p = doc.paragraphs[7]
            p.text = f"ИТОГО:  {str(z1).replace('.', ',')} руб."
            table = doc.tables[4]
            k = table.cell(0, 0)
            k.text = f'Особые отметки: ' + number[4]
            doc.save('of.docx')
        else:
            doc = Document('sch_list.docx')
            table = doc.tables[0]
            cell = table.cell(0, 0)
            new_text = cell.text.split('\n')
            new_text[1] += str(number[0])
            new_text[2] += number[6]
            new_text[3] += number[9]
            new_text[5] += number[3]
            new_text[6] += number[8]
            new_text[7] += number[2]
            cell.text = '\n'.join(new_text)
            table = doc.tables[1]
            row = table.rows[-1]
            ss = load_services1(number)
            for i in range(6):
                self.remove_row(table, table.rows[1])
            o = 0
            z = 0
            for i in range(len(ss)):
                if ss[i][3] != '0':
                    table.add_row()
                    k = table.cell(i + 2 - o, 0)
                    k.text = str(i + 1 - o)
                    k = table.cell(i + 2 - o, 1)
                    k.text = ss[i][1]
                    k = table.cell(i + 2 - o, 3)
                    k.text = ss[i][3]
                    k = table.cell(i + 2 - o, 4)
                    k.text = str(float(ss[i][2].replace(',', '.')) / int(ss[i][3])).replace('.', ',')
                    k = table.cell(i + 2 - o, 5)
                    k.text = ss[i][2]
                    z += float(ss[i][2].replace(',', '.'))
                else:
                    o += 1
            k = table.cell(1, 5)
            k.text = str(z).replace('.', ',')
            row0 = table.rows[1]
            row1 = table.rows[-1]
            row1._tr.addnext(row0._tr)
            p = doc.paragraphs[5]

            table = doc.tables[2]
            ss = load_sales(number)
            for i in range(4):
                self.remove_row(table, table.rows[1])
            z1 = z
            z = 0
            o = 0
            for i in range(len(ss)):
                if ss[i][3] != '0':
                    table.add_row()
                    k = table.cell(i + 2 - o, -1)
                    k.text = str(i + 1 - o)
                    k = table.cell(i + 2 - o, 0)
                    k.text = ss[i][1]
                    k = table.cell(i + 2 - o, 2)
                    k.text = ss[i][3]
                    k = table.cell(i + 2 - o, 3)
                    k.text = str(float(ss[i][2].replace(',', '.')) / int(ss[i][3])).replace('.', ',')
                    k = table.cell(i + 2 - o, 4)
                    k.text = ss[i][2]
                    k = table.cell(i + 2 - o, 5)
                    k.text = '0'
                    z += float(ss[i][2].replace(',', '.'))
                else:
                    o += 1
            k = table.cell(1, 5)
            k.text = str(z).replace('.', ',')
            z1 += z
            row0 = table.rows[1]
            row1 = table.rows[-1]
            row1._tr.addnext(row0._tr)
            p = doc.paragraphs[6]
            p.text = f"ИТОГО к оплате:  {str(z1).replace('.', ',')} руб."
            table = doc.tables[3]
            k = table.cell(0, 0)
            k.text = f'Особые отметки: ' + number[4]
            doc.save('of.docx')
        os.startfile('of.docx', 'open')


    def rec(self):
        windows.append(Rec(self.number, self))
        windows[-1].show()

    def new_sale(self):
        windows.append(New_sale(self.g, self, self.number[0]))
        windows[-1].show()

    def password_editing(self):
        len_1 = len(self.ui.l_password.text())
        len_2 = len(self.passwo)
        if len_1 < len_2:
            self.passwo = self.passwo[:(len_1-len_2)]
        elif len_1 > len_2:
            self.passwo = self.passwo + self.ui.l_password.text()[-1]
        self.ui.l_password.setText('*' * len_1)

    def update_sost(self):
        self.sost_to_save = ''
        if self.ui.btn_sost.text() == 'Новый':
            self.ui.btn_sost.setText('В работе')
            self.ui.la_password.hide()
            self.ui.l_password.hide()
            self.ui.pushButton.hide()
            self.sost_to_save = 'in'
        elif self.ui.btn_sost.text() == 'В работе':
            self.ui.btn_sost.setText('В работе, ждёт запчастей')
            self.sost_to_save = 'wa'
        elif self.ui.btn_sost.text() == 'В работе, ждёт запчастей':
            self.ui.btn_sost.setText('Готов')
            self.ui.la_password.hide()
            self.ui.l_password.hide()
            self.ui.pushButton.hide()
            self.sost_to_save = 'ne'
        elif self.ui.btn_sost.text() == 'Готов':
            o = load_services1(self.number)
            k = 1
            for i in o:
                if not i[6]:
                    if i[3] != '0':
                        k = 0
            if k:
                self.ui.btn_sost.setText('Закрыт')
                self.ui.la_password.hide()
                self.ui.l_password.hide()
                self.ui.pushButton.hide()
                self.ui.btn_close_offer.show()
                self.ui.l_password.show()
                self.ui.la_password.show()
                self.sost_to_save = 'co'
            else:
                self.ui.btn_close_offer.hide()
                self.ui.l_password.hide()
                self.ui.la_password.hide()
                self.ui.btn_sost.setText('Отменён')
                self.sost_to_save = 'cl'
                self.ui.la_password.show()
                self.ui.l_password.show()
                self.ui.pushButton.show()
        elif self.ui.btn_sost.text() == 'Закрыт':
            self.ui.btn_close_offer.hide()
            self.ui.l_password.hide()
            self.ui.la_password.hide()
            self.ui.btn_sost.setText('Отменён')
            self.sost_to_save = 'cl'
            self.ui.la_password.show()
            self.ui.l_password.show()
            self.ui.pushButton.show()
        elif self.ui.btn_sost.text() == 'Отменён':
            self.ui.la_password.hide()
            self.ui.l_password.hide()
            self.ui.pushButton.hide()
            self.ui.btn_sost.setText('Новый')
            self.sost_to_save = 'nw'

    def delete_off(self):
        if self.passwo == str(load_pass()[0][0]):
            delete_offer(self.number[0])
            self.g.offers_open()
            self.close()


    def close_realy(self):
        if self.passwo == str(load_pass()[0][0]):
            d_now = datetime.now().date()
            dd = f'{d_now.year}-{d_now.month}-{d_now.day}'
            self.number = list(self.number)
            self.number[9] = dd
            self.sost_to_save = 'cor'
            update_sost(str(self.number[0]), 'cor', date=dd)
            self.ui.l_date_1.setText(dd)
            self.ui.l_date_1.show()
            self.ui.label_9.show()
            self.ui.btn_add_service.hide()
            self.ui.btn_sost.setEnabled(False)
            self.ui.btn_add_service.hide()
            self.ui.btn_close_offer.hide()
            self.ui.l_password.hide()
            self.ui.la_password.hide()
            self.ui.btn_add_sale.hide()
            self.ui.l_client_.setEnabled(False)
            self.ui.l_type.setEnabled(False)
            self.ui.l_numberr.setEnabled(False)
            self.g.offers_open()

    def update_table(self):
        c = 0
        r = self.ui.tableWidget.rowCount()
        for i in range(r):
            self.ui.tableWidget.removeRow(0)
        self.services = load_services1(self.number)
        o = self.services
        self.ui.tableWidget.setRowCount(len(o))
        z = len(o)
        for i in range(len(o)):
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(o[i][4]))
            if int(o[i][3]) != 0:
                self.ui.tableWidget.setItem(i, 2, QtWidgets.QTableWidgetItem(str(float(o[i][2].replace(',', '.')) / int(o[i][3])).replace('.', ',')))
            else:
                self.ui.tableWidget.setItem(i, 2, QtWidgets.QTableWidgetItem('0'))
            self.ui.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(o[i][1]))
            self.ui.tableWidget.setItem(i, 3, QtWidgets.QTableWidgetItem(o[i][3]))
            self.ui.tableWidget.setItem(i, 4, QtWidgets.QTableWidgetItem(o[i][2]))
            self.ui.tableWidget.setItem(i, 7, QtWidgets.QTableWidgetItem(o[i][7]))
            c += float(o[i][2].replace(',', '.'))
            if o[i][6] == 1:
                btn = QtWidgets.QPushButton('Готово')
                btn.clicked.connect(self.stat_ch)
                btn.offer_number = self.number
                btn.serv_info = o[i]
                self.ui.tableWidget.setCellWidget(i, 5, btn)
            else:
                btn = QtWidgets.QPushButton('В работе')
                btn.clicked.connect(self.stat_ch)
                btn.offer_number = self.number
                btn.serv_info = o[i]
                self.ui.tableWidget.setCellWidget(i, 5, btn)
            self.ui.tableWidget.setItem(i, 6, QtWidgets.QTableWidgetItem(o[i][5].split(' ')[0]))
            btn.setEnabled(self.closeee)
        sales = load_sales(self.number)
        self.ui.tableWidget.setRowCount(len(sales) + len(o))
        for i in range(len(sales)):
            self.ui.tableWidget.setItem(i + z, 0, QtWidgets.QTableWidgetItem(sales[i][4]))
            if int(sales[i][3]) != 0:
                self.ui.tableWidget.setItem(i + z, 2, QtWidgets.QTableWidgetItem(
                str(float(sales[i][2].replace(',', '.')) / int(sales[i][3])).replace('.', ',')))
            else:
                self.ui.tableWidget.setItem(i + z, 2, QtWidgets.QTableWidgetItem('0'))
            self.ui.tableWidget.setItem(i + z, 1, QtWidgets.QTableWidgetItem('Продано: ' + sales[i][1]))
            self.ui.tableWidget.setItem(i + z, 3, QtWidgets.QTableWidgetItem(sales[i][3]))
            self.ui.tableWidget.setItem(i + z, 4, QtWidgets.QTableWidgetItem(sales[i][2]))
            c += float(sales[i][2].replace(',', '.'))
        self.ui.l_ito.setText('ИТОГО:' + ' ' + str(c))
        self.ui.tableWidget.resizeColumnsToContents()

    def stat_ch(self):
        windows.append(Stat_change(self))
        windows[-1].show()

    def new_service(self):
        windows.append(New_service(self.g, self, self.number[0]))
        windows[-1].show()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)


class More_service(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        self.number = QtWidgets.QApplication.instance().sender()
        super(More_service, self).__init__()
        self.info = list(filter(lambda s: s[0] == self.number.number_of_service, load_services()))[0]
        self.ui = Ui_more_service()
        self.ui.setupUi(self)
        self.ui.bt_add.clicked.connect(self.confirm)
        self.ui.line_sur.setText(self.info[1])
        pr = self.info[2].replace(',', '.')
        pr2 = self.info[4].replace(',', '.')
        group = self.info[5]
        self.ui.db_price.setValue(float(pr))
        self.ui.bt_close.clicked.connect(self.close)
        self.ui.db_price_2.setValue(float(pr2))
        self.ui.bt_delete.clicked.connect(self.delete)
        self.ui.comboBox.addItem(group)
        self.ui.comboBox.addItem('Все')
        k = load_groups()
        for i in k:
            if i[0] != group:
                self.ui.comboBox.addItem(i[0])

    def confirm(self):
        if self.ui.line_sur.text().strip() != '':
            update_services(self.number.number_of_service, self.ui.line_sur.text(), self.ui.db_price.text(), self.ui.db_price_2.text(), self.ui.comboBox.currentText().strip())
            self.g.services_open()
            self.close()

    def delete(self):
        delete_service(self.number.number_of_service)
        self.g.services_open()
        self.close()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)
        print(windows)


class Add_service(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        super(Add_service, self).__init__()
        self.ui = Ui_add_service()
        self.ui.setupUi(self)
        self.ui.bt_add.clicked.connect(self.confirm)
        self.ui.bt_close.clicked.connect(self.close)
        k = load_groups()
        for i in k:
            self.ui.comboBox.addItem(i[0])

    def confirm(self):
        if self.ui.line_sur.text().strip() != '':
            save_service(self.ui.line_sur.text().strip(), self.ui.db_price.text(), self.ui.db_price_2.text(), self.ui.comboBox.currentText().strip())
            self.g.services_open()
            self.close()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)
        print(windows)


class Add_sale(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        super(Add_sale, self).__init__()
        self.ui = Ui_add_sale()
        self.ui.setupUi(self)
        self.ui.bt_add.clicked.connect(self.confirm)
        self.ui.bt_close.clicked.connect(self.close)

    def confirm(self):
        if self.ui.line_sur.text().strip() != '':
            save_sale(self.ui.line_sur.text().strip(), self.ui.db_price.text(), self.ui.s_col.text())
            self.g.in_open()
            self.close()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)
        print(windows)


class More_sale(QtWidgets.QMainWindow):
    def __init__(self, g, number):
        self.number = number
        self.g = g
        s = list(filter(lambda s: str(s[0]) == str(number), load_in()))[0]
        super(More_sale, self).__init__()
        self.ui = Ui_more_sale()
        self.ui.setupUi(self)
        self.ui.line_sur.setText(s[1])
        self.ui.db_price.setValue(float(s[3].replace(',', '.')))
        self.ui.s_col.setValue(int(s[2]))
        self.ui.bt_add.clicked.connect(self.confirm)
        self.ui.bt_close.clicked.connect(self.close)

    def confirm(self):
        if self.ui.line_sur.text().strip() != '':
            update_sale_1(self.number, self.ui.line_sur.text().strip(), self.ui.db_price.text(), self.ui.s_col.text())
            self.g.in_open()
            self.close()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)
        print(windows)


class Add_offer(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        super(Add_offer, self).__init__()
        self.ui = Ui_add_offer()
        self.ui.setupUi(self)
        self.ui.l_date.setText(self.ui.l_date.text() + ' ' + str(datetime.now().date()))
        self.ui.btn_confirm.clicked.connect(self.confirm)
        self.ui.btn_cancel.clicked.connect(self.close)

    def confirm(self):
        try:
            if self.ui.line_of_name.text().strip() != '' and self.ui.line_of_client.text().strip() != '' and ' ' not in self.ui.line_of_name.text().strip() and self.ui.comboBox.currentText() !='':
                date = str(datetime.now().date())
                type = ''
                if self.ui.comboBox.currentText() == 'Легковой':
                    type = 'l'
                else:
                    type = 'g'
                save_offer(name=self.ui.line_of_name.text().strip().capitalize(), client=self.ui.line_of_client.text().strip(), date=date, phone=self.ui.l_number.text(), worker=type)
                offers = list(filter(lambda s: s[2].lower().strip() == self.ui.line_of_name.text().lower().strip(), load_offers()))[-1]
                self.g.offers_open()
                windows.append(More_offer(self.g, offers))
                windows[-1].show()
                self.close()
            else:
                raise sqlite3.OperationalError
        except sqlite3.OperationalError:
            self.ui.l_exception.setText("Присутствуют недопустимые символы.")

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)
        print(windows)


class More_workers(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        self.worker_status = list(filter(lambda s: s[1] == g, load_workers()))[0][3]
        proc = int(list(filter(lambda s: s[1] == g, load_workers()))[0][2])
        self.passwo = ''
        self.worker_info = load_worker_service(g)
        super(More_workers, self).__init__()
        self.ui = UI_more_workers()
        self.ui.setupUi(self)
        self.ui.s_procent.setValue(proc)
        self.ui.label_2.setText(g.replace('ALEEE', ' '))
        self.ui.tableWidget.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self.ui.b_use.clicked.connect(self.confirm)
        self.ui.btn_change.clicked.connect(self.change)
        self.ui.btn_save.clicked.connect(self.save)
        self.ui.l_password.textEdited.connect(self.password_editing)
        self.ui.s_procent.hide()
        self.ui.btn_save.hide()
        self.ui.btn_fired.clicked.connect(self.fired)
        if not self.worker_status:
            self.ui.btn_fired.setText('Восстановить')
        now_date = datetime.now()
        s_date = datetime(now_date.year, now_date.month, 1).date()
        n_date = now_date.date()
        day = str(s_date.day)
        if len(day) == 1:
            day = '0' + day
        month = str(s_date.month)
        if len(month) == 1:
            month = '0' + month
        self.ui.l_syear.setText(str(s_date.year))
        self.ui.l_sday.setText(str(day))
        self.ui.l_smonth.setText(str(month))
        day = str(n_date.day)
        if len(day) == 1:
            day = '0' + day
        month = str(n_date.month)
        if len(month) == 1:
            month = '0' + month
        self.ui.l_pyear.setText(str(n_date.year))
        self.ui.l_pday.setText(str(day))
        self.ui.l_pmonth.setText(str(month))
        o = 0
        sum = 0
        comp = []
        if self.worker_info != []:
            self.ui.tableWidget.setRowCount(len(self.worker_info))
            for i in range(len(self.worker_info)):
                procent = self.worker_info[i][-1]
                datt = self.worker_info[i][3].split('-')
                off = self.worker_info[i][0]
                datt = datetime(year=int(datt[0]), month=int(datt[1]), day=int(datt[2])).date()
                offee = list(filter(lambda s: str(s[0]) == str(off), load_offers()))[0]
                sost = offee[7]
                passes = 0
                if sost == 'cor' and offee[9] != None:
                    datt = offee[9].split('-')
                    datt = datetime(year=int(datt[0]), month=int(datt[1]), day=int(datt[2])).date()
                    if s_date <= datt <= n_date:
                        passes = 1
                elif sost == 'cor':
                    if s_date <= datt <= n_date:
                        passes = 1
                if passes:
                    self.ui.tableWidget.setItem(i - o, 0, QtWidgets.QTableWidgetItem(self.worker_info[i][3]))
                    bt = QtWidgets.QPushButton('№' + str(self.worker_info[i][0]))
                    bt.offer_number = int(self.worker_info[i][0])
                    bt.clicked.connect(self.offer_open)
                    self.ui.tableWidget.setCellWidget(i - o, 1, bt)
                    self.ui.tableWidget.setItem(i - o, 4, QtWidgets.QTableWidgetItem(
                        str(float(self.worker_info[i][1].replace(',', '.')) * float(procent) / 100).replace('.', ',')))
                    sum += float(self.worker_info[i][1].replace(',', '.')) * float(procent) / 100
                    self.ui.tableWidget.setItem(i - o, 2, QtWidgets.QTableWidgetItem(self.worker_info[i][2]))
                    self.ui.tableWidget.setItem(i - o, 3, QtWidgets.QTableWidgetItem(self.worker_info[i][1]))
                else:
                    o += 1
        self.ui.l_ito.setText('Зарплата за указанный период: ' + str(sum).replace('.', ','))
        self.ui.tableWidget.setRowCount(len(self.worker_info) - o)

    def fired(self):
        if self.worker_status:
            fire(self.g, 0)
            self.ui.btn_fired.setText('Восстановить')
            self.worker_status = 0
        else:
            fire(self.g, 1)
            self.ui.btn_fired.setText('Уволить')
            self.worker_status = 1

    def save(self):
        update_procent(self.g, self.ui.s_procent.value())
        self.ui.s_procent.hide()
        self.ui.btn_save.hide()
        self.passwo = ''
        self.ui.l_password.setText('')

    def change(self):
        if self.passwo == str(load_pass()[0][0]):
            self.ui.s_procent.show()
            self.ui.btn_save.show()

    def confirm(self):
        self.worker_info = load_worker_service(self.g)
        s_date = datetime(int(self.ui.l_syear.text()), int(self.ui.l_smonth.text()), int(self.ui.l_sday.text())).date()
        n_date = datetime(int(self.ui.l_pyear.text()), int(self.ui.l_pmonth.text()), int(self.ui.l_pday.text())).date()
        o = 0
        sum = 0
        for i in range(self.ui.tableWidget.rowCount()):
            self.ui.tableWidget.removeRow(0)
        o = 0
        sum = 0
        if self.worker_info != []:
            self.ui.tableWidget.setRowCount(len(self.worker_info))
            for i in range(len(self.worker_info)):
                procent = self.worker_info[i][-1]
                datt = self.worker_info[i][3].split('-')
                off = self.worker_info[i][0]
                datt = datetime(year=int(datt[0]), month=int(datt[1]), day=int(datt[2])).date()
                offee = list(filter(lambda s: str(s[0]) == str(off), load_offers()))[0]
                sost = offee[7]
                passes = 0
                if sost == 'cor' and offee[9] != None:
                    datt = offee[9].split('-')
                    datt = datetime(year=int(datt[0]), month=int(datt[1]), day=int(datt[2])).date()
                    if s_date <= datt <= n_date:
                        passes = 1
                elif sost == 'cor':
                    if s_date <= datt <= n_date:
                        passes = 1
                if passes:
                    self.ui.tableWidget.setItem(i - o, 0, QtWidgets.QTableWidgetItem(self.worker_info[i][3]))
                    bt = QtWidgets.QPushButton('№' + str(self.worker_info[i][0]))
                    bt.offer_number = int(self.worker_info[i][0])
                    bt.clicked.connect(self.offer_open)
                    self.ui.tableWidget.setCellWidget(i - o, 1, bt)
                    self.ui.tableWidget.setItem(i - o, 4, QtWidgets.QTableWidgetItem(
                        str(float(self.worker_info[i][1].replace(',', '.')) * float(procent) / 100).replace('.', ',')))
                    sum += float(self.worker_info[i][1].replace(',', '.')) * float(procent) / 100
                    self.ui.tableWidget.setItem(i - o, 2, QtWidgets.QTableWidgetItem(self.worker_info[i][2]))
                    self.ui.tableWidget.setItem(i - o, 3, QtWidgets.QTableWidgetItem(self.worker_info[i][1]))
                else:
                    o += 1
        self.ui.l_ito.setText('Зарплата за указанный период: ' + str(sum).replace('.', ','))
        self.ui.tableWidget.setRowCount(len(self.worker_info) - o)

    def password_editing(self):
        len_1 = len(self.ui.l_password.text())
        len_2 = len(self.passwo)
        if len_1 < len_2:
            self.passwo = self.passwo[:(len_1-len_2)]
        elif len_1 > len_2:
            self.passwo = self.passwo + self.ui.l_password.text()[-1]
        self.ui.l_password.setText('*' * len_1)

    def offer_open(self):
        btn = QtWidgets.QApplication.instance().sender()
        number = btn.offer_number
        offer = list(filter(lambda s: s[0] == number, load_offers()))[0]
        windows.append(More_offer(self.g, offer))
        windows[-1].show()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)


class Add_worker(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        super(Add_worker, self).__init__()
        self.ui = add_workerses()
        self.ui.setupUi(self)
        self.ui.bt_add.clicked.connect(self.nice)
        self.ui.bt_close.clicked.connect(self.nein)

    def closeEvent(self, a0: QtGui.QCloseEvent):
        self.close()
        windows.remove(self)

    def nice(self):
        try:
            if self.ui.line_name.text().strip() != '' and self.ui.line_sur.text().strip() != '' and self.ui.line_ot.text().capitalize().strip() != '' and ';' not in self.ui.line_name.text().strip() + self.ui.line_sur.text().strip() + self.ui.line_ot.text().capitalize().strip():

                save_worker(f'{self.ui.line_sur.text().capitalize().strip()}ALEEE{self.ui.line_name.text().capitalize().strip()}ALEEE{self.ui.line_ot.text().capitalize().strip()}', procent=self.ui.s_procent.text())
                self.g.workers_open()
                self.close()
            else:
                raise sqlite3.OperationalError
        except sqlite3.OperationalError:
            self.ui.l_exception.setText("Присутствуют недопустимые символы.")
    def nein(self):
        self.close()


class Add_rash(QtWidgets.QMainWindow):
    def __init__(self, g):
        self.g = g
        super(Add_rash, self).__init__()
        self.ui = Ui_add_rash()
        self.ui.setupUi(self)
        self.ui.bt_add.clicked.connect(self.confirm)
        self.ui.bt_close.clicked.connect(self.close)

    def confirm(self):
        if self.ui.l_name.text().strip() != '':
            save_rash(self.ui.l_name.text(), self.ui.s_col.text(), self.ui.s_price.text(), datetime.now().date())
            self.g.ras_open()
            self.close()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        windows.remove(self)
        print(windows)


class More_rash(QtWidgets.QMainWindow):
    def __init__(self, g, number):
        self.g = g
        self.number = number
        super(More_rash, self).__init__()
        self.ui = Ui_more_rash()
        self.ui.setupUi(self)
        self.ui.bt_add.clicked.connect(self.confirm)
        self.ui.bt_close.clicked.connect(self.close)
        o = list(filter(lambda s: s[0] == number, load_ras()))[0]
        self.ui.l_name.setText(o[1])
        self.ui.s_col.setValue(int(o[2]))
        self.ui.s_price.setValue(float(o[3].replace(',', '.')))
        self.ui.bt_add_2.clicked.connect(self.delete)

    def confirm(self):
        if self.ui.l_name.text().strip() != '':
            update_ras(self.number, self.ui.s_col.text(), self.ui.s_price.text(), self.ui.l_name.text().strip())
        self.g.ras_open()
        self.close()

    def delete(self):
        delete_ras(self.number)
        self.g.ras_open()
        self.close()

    def closeEvent(self, a0: QtGui.QCloseEvent):
        windows.remove(self)
        print(windows)


class mywindow(QtWidgets.QMainWindow):
    def __init__(self):
        self.state = 'off'
        super(mywindow, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.tableWidget.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self.ui.bt_workers.clicked.connect(self.workers_open)
        self.ui.bt_offers.clicked.connect(self.offers_open)
        self.ui.tableWidget_2.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self.ui.bt_add_worker.clicked.connect(self.add_workerss)
        self.ui.bt_add.clicked.connect(self.add_offer)
        self.ui.bt_services.clicked.connect(self.services_open)
        self.offers_open()
        self.ui.bt_add_service.clicked.connect(self.add_services)
        self.ui.b_find.clicked.connect(self.number_find)
        self.ui.bt_used.clicked.connect(self.ras_open)
        self.ui.pushButton.clicked.connect(self.add_rashod)
        now_date = datetime.now()
        s_date = datetime(now_date.year, now_date.month, 1).date()
        self.s_date = s_date
        self.s_date_2 = s_date
        n_date = now_date.date()
        self.n_date = n_date
        self.n_date_2 = n_date
        day = str(s_date.day)
        if len(day) == 1:
            day = '0' + day
        month = str(s_date.month)
        if len(month) == 1:
            month = '0' + month
        self.ui.l_syear.setText(str(s_date.year))
        self.ui.l_sday.setText(str(day))
        self.ui.l_smonth.setText(str(month))
        self.ui.l_syear_2.setText(str(s_date.year))
        self.ui.l_sday_2.setText(str(day))
        self.ui.l_smonth_2.setText(str(month))
        day = str(n_date.day)
        if len(day) == 1:
            day = '0' + day
        month = str(n_date.month)
        if len(month) == 1:
            month = '0' + month
        self.ui.l_pyear.setText(str(n_date.year))
        self.ui.l_pday.setText(str(day))
        self.ui.l_pmonth.setText(str(month))
        self.ui.l_pyear_2.setText(str(n_date.year))
        self.ui.l_pday_2.setText(str(day))
        self.ui.l_pmonth_2.setText(str(month))
        self.ui.btn_use.clicked.connect(self.ras_open)
        self.ui.bt_stat.clicked.connect(self.stat_open)
        self.stat = self.services_stat
        self.ui.bt_workers_stat.clicked.connect(self.worker_stat)
        self.ui.bt_services_stat.clicked.connect(self.services_stat)
        self.ui.btn_use_2.clicked.connect(self.stat_open)
        self.ui.bt_offers_stat.clicked.connect(self.ob_stat)
        self.ui.bt_have.clicked.connect(self.in_open)
        self.ui.pushButton_2.clicked.connect(self.add_sale)
        self.ui.s_page.valueChanged.connect(self.update_page)
        self.ui.add_group.clicked.connect(self.add_group)
        self.ui.service_group.currentTextChanged.connect(self.services_open)
        self.ui.pushButton_3.clicked.connect(self.delete_gp)
        gr = load_groups()
        for i in gr:
            self.ui.service_group.addItem(i[0])

    def delete_gp(self):
        delete_group(self.ui.service_group.currentText())
        self.ui.service_group.clear()
        self.ui.service_group.addItem('Все')
        gr = load_groups()
        for i in gr:
            self.ui.service_group.addItem(i[0])
        self.services_open()

    def wheelEvent(self, event):
        if self.state == 'off':
            angle = event.angleDelta()
            angleX = angle.x()
            angleY=angle.y()
            if angleY == 120:
                self.ui.s_page.setValue(int(self.ui.s_page.text()) + 1)
            elif angleY == -120:
                self.ui.s_page.setValue(int(self.ui.s_page.text()) - 1)

    def add_sale(self):
        windows.append(Add_sale(self))
        windows[-1].show()

    def add_group(self):
        windows.append(Add_group(self))
        windows[-1].show()

    def stat_open(self):
        self.state = 'stat'
        self.ui.gp_offer.hide()
        self.ui.groupBox.hide()
        self.ui.gp_services.hide()
        self.ui.gp_ras.hide()
        self.ui.gp_in.hide()
        self.ui.gp_stat.show()
        self.stat()

    def in_open(self):
        self.state = 'in'
        for i in range(self.ui.tableWidget_8.rowCount()):
            self.ui.tableWidget_8.removeRow(0)
        self.ui.gp_stat.hide()
        self.ui.gp_offer.hide()
        self.ui.groupBox.hide()
        self.ui.gp_services.hide()
        self.ui.gp_ras.hide()
        self.ui.gp_in.show()
        inn = load_in()
        inn = sorted(inn, key=lambda s: s[1])
        self.ui.tableWidget_8.setRowCount(len(inn))
        for i in range(len(inn)):
            self.ui.tableWidget_8.setItem(i, 0, QtWidgets.QTableWidgetItem(inn[i][1]))
            self.ui.tableWidget_8.setItem(i, 1, QtWidgets.QTableWidgetItem(str(inn[i][2])))
            btn = QtWidgets.QPushButton('Добавить')
            btn.in_number = inn[i][0]
            btn.clicked.connect(self.plus_in)
            self.ui.tableWidget_8.setCellWidget(i, 2, btn)
            btn = QtWidgets.QPushButton('Убавить')
            btn.in_number = inn[i][0]
            btn.clicked.connect(self.minus_in)
            self.ui.tableWidget_8.setCellWidget(i, 3, btn)
            btn = QtWidgets.QPushButton('Изменить')
            btn.in_number = inn[i][0]
            btn.clicked.connect(self.izm_in)
            self.ui.tableWidget_8.setCellWidget(i, 4, btn)
        self.ui.tableWidget_8.resizeColumnsToContents()

    def izm_in(self):
        num = QtWidgets.QApplication.instance().sender().in_number
        windows.append(More_sale(self, num))
        windows[-1].show()

    def minus_in(self):
        num = QtWidgets.QApplication.instance().sender().in_number
        inn = list(filter(lambda s: s[0] == num, load_in()))[0][2]
        update_in(num, inn - 1)
        self.in_open()

    def plus_in(self):
        num = QtWidgets.QApplication.instance().sender().in_number
        inn = list(filter(lambda s: s[0] == num, load_in()))[0][2]
        update_in(num, inn + 1)
        self.in_open()

    def ob_stat(self):
        for i in range(self.ui.tableWidget_7.rowCount()):
            self.ui.tableWidget_7.removeRow(0)
        self.ui.gb_stat_service.hide()
        self.ui.gb_stat_workers.hide()
        self.ui.gb_ob_stat.show()
        self.stat = self.ob_stat
        self.s_date_2 = datetime(int(self.ui.l_syear_2.text()), int(self.ui.l_smonth_2.text()),
                                 int(self.ui.l_sday_2.text())).date()
        self.n_date_2 = datetime(int(self.ui.l_pyear_2.text()), int(self.ui.l_pmonth_2.text()),
                                 int(self.ui.l_pday_2.text())).date()
        rass = []
        rs = load_ras()
        for i in rs:
            date = i[4].split('-')
            date = datetime(int(date[0]), int(date[1]), int(date[2])).date()
            if self.s_date_2 <= date <= self.n_date_2:
                rass.append(i)
        off = []
        offers = load_offers()
        optimus = {}
        for i in offers:
            s = [str(i[0]), 0]
            optimus[i[0]] = list(i).copy()
            if i[7] == 'cor':
                date = i[9].split('-')
                datt = datetime(year=int(date[0]), month=int(date[1]), day=int(date[2])).date()
                if self.s_date_2 <= datt <= self.n_date_2:
                    services = load_services1(i)
                    for j in services:
                        date = j[4].split('-')
                        datt = datetime(year=int(date[0]), month=int(date[1]), day=int(date[2])).date()
                        offee = optimus[i[0]].copy()
                        sost = offee[7]
                        passes = 0
                        date = datetime(int(date[0]), int(date[1]), int(date[2])).date()
                        if sost == 'cor' and offee[9] != None:
                            datt = offee[9].split('-')
                            date = datetime(year=int(datt[0]), month=int(datt[1]), day=int(datt[2])).date()
                            if self.s_date_2 <= date <= self.n_date_2:
                                passes = 1
                        elif sost == 'cor':
                            if self.s_date_2 <= date <= self.n_date_2:
                                passes = 1
                        if passes:
                            s[1] += float(j[2].replace(',', '.'))
                    o = load_sales(i)
                    for io in range(len(o)):
                        date = o[io][4].split('-')
                        date = datetime(int(date[0]), int(date[1]), int(date[2])).date()
                        if self.s_date_2 <= date <= self.n_date_2:
                            s[1] += float(o[io][2].replace(',', '.'))
                    if s[-1] != 0:
                        off.append(s)
        self.ui.tableWidget_7.setRowCount(len(off) + len(rass))
        sum = 0
        for i in range(len(off)):
            self.ui.tableWidget_7.setItem(i, 0, QtWidgets.QTableWidgetItem('Заказ'))
            self.ui.tableWidget_7.setItem(i, 1, QtWidgets.QTableWidgetItem(str(off[i][0])))
            self.ui.tableWidget_7.setItem(i, 2, QtWidgets.QTableWidgetItem(str(off[i][1]).replace('.', ',')))
            sum += off[i][1]
        for i in range(len(rass)):
            self.ui.tableWidget_7.setItem(i + len(off), 0, QtWidgets.QTableWidgetItem('Расход'))
            self.ui.tableWidget_7.setItem(i + len(off), 1, QtWidgets.QTableWidgetItem(str(rass[i][1])))
            self.ui.tableWidget_7.setItem(i + len(off), 2, QtWidgets.QTableWidgetItem(str(int(rass[i][2]) * float(str(rass[i][3]).replace(',', '.'))).replace('.', ',')))
            sum -= int(rass[i][2]) * float(str(rass[i][3]).replace(',', '.'))
        self.ui.label_7.setText('Прибыль: ' + str(sum).replace('.', ','))

    def worker_stat(self):
        for i in range(self.ui.tableWidget_6.rowCount()):
            self.ui.tableWidget_6.removeRow(0)
        self.ui.gb_ob_stat.hide()
        self.ui.gb_stat_service.hide()
        self.ui.gb_stat_workers.show()
        self.stat = self.worker_stat
        self.s_date_2 = datetime(int(self.ui.l_syear_2.text()), int(self.ui.l_smonth_2.text()),
                                 int(self.ui.l_sday_2.text())).date()
        self.n_date_2 = datetime(int(self.ui.l_pyear_2.text()), int(self.ui.l_pmonth_2.text()),
                                 int(self.ui.l_pday_2.text())).date()
        k = {}
        workers = load_workers()
        for i in workers:
            k[' '.join(i[1].split('ALEEE'))] = 0
        sum = 0
        optimus = {}
        for i in workers:
            services = load_worker_service(i[1])
            name = ' '.join(i[1].split('ALEEE'))
            for j in services:
                procent = int(j[-1])
                date = j[3].split('-')
                off = j[0]
                datt = datetime(year=int(date[0]), month=int(date[1]), day=int(date[2])).date()
                if off not in optimus:
                    offee = list(filter(lambda s: str(s[0]) == str(off), load_offers()))[0]
                    optimus[off] = list(offee).copy()
                else:
                    offee = optimus[off].copy()
                sost = offee[7]
                passes = 0
                date = datetime(int(date[0]), int(date[1]), int(date[2])).date()
                if sost == 'cor' and offee[9] != None:
                    datt = offee[9].split('-')
                    date = datetime(year=int(datt[0]), month=int(datt[1]), day=int(datt[2])).date()
                    if self.s_date_2 <= date <= self.n_date_2:
                        passes = 1
                elif sost == 'cor':
                    if self.s_date_2 <= date <= self.n_date_2:
                        passes = 1
                if passes:
                    k[name] += (procent * float(j[1].replace(',', '.')) / 100)
                    sum += (procent * float(j[1].replace(',', '.')) / 100)
        self.ui.label_6.setText('Итого по зарплате: ' + str(sum).replace('.', ','))
        kee = list(k.keys())
        kee = sorted(kee)
        zxcvb = len(kee)
        self.ui.tableWidget_6.setRowCount(len(kee))
        hj = 0
        for i in range(len(kee)):
            if k[kee[i]] != 0:
                self.ui.tableWidget_6.setItem(i - hj, 0, QtWidgets.QTableWidgetItem(str(kee[i])))
                self.ui.tableWidget_6.setItem(i - hj, 1, QtWidgets.QTableWidgetItem(str(k[kee[i]]).replace('.', ',')))
            else:
                hj += 1
                zxcvb -= 1
        self.ui.tableWidget_6.setRowCount(zxcvb)

    def services_stat(self):
        self.ui.gb_ob_stat.hide()
        self.ui.gb_stat_workers.hide()
        self.ui.gb_stat_service.show()
        self.stat = self.services_stat
        for i in range(self.ui.tableWidget_5.rowCount()):
            self.ui.tableWidget_5.removeRow(0)
        self.s_date_2 = datetime(int(self.ui.l_syear_2.text()), int(self.ui.l_smonth_2.text()),
                               int(self.ui.l_sday_2.text())).date()
        self.n_date_2 = datetime(int(self.ui.l_pyear_2.text()), int(self.ui.l_pmonth_2.text()),
                               int(self.ui.l_pday_2.text())).date()
        offers = load_offers()
        s = load_services()
        k = {}
        for i in s:
            k[i[1]] = [0, 0]
        optimus = {}
        for i in range(len(offers)):
            optimus[offers[i][0]] = list(offers[i]).copy()
            if offers[i][7] == 'cor':
                date = offers[i][9].split('-')
                datt = datetime(year=int(date[0]), month=int(date[1]), day=int(date[2])).date()
                if self.s_date_2 <= datt <= self.n_date_2:
                    services = load_services1(offers[i])
                    for j in services:
                        date = j[4].split('-')
                        datt = datetime(year=int(date[0]), month=int(date[1]), day=int(date[2])).date()
                        offee = optimus[offers[i][0]]
                        sost = offee[7]
                        passes = 0
                        date = datetime(int(date[0]), int(date[1]), int(date[2])).date()
                        if sost == 'cor' and offee[9] != None:
                            datt = offee[9].split('-')
                            date = datetime(year=int(datt[0]), month=int(datt[1]), day=int(datt[2])).date()
                            if self.s_date_2 <= date <= self.n_date_2:
                                passes = 1
                        elif sost == 'cor':
                            if self.s_date_2 <= date <= self.n_date_2:
                                passes = 1
                        if passes:
                            if j[1] in k:
                                k[j[1]][0] += int(j[3])
                                k[j[1]][1] += float(j[2].replace(',', '.'))
                            else:
                                k[j[1]] = [0, 0]
                                k[j[1]][0] = int(j[3])
                                k[j[1]][1] = float(j[2].replace(',', '.'))
        kee = list(k.keys())
        kee = sorted(kee, key=lambda s: k[s][1], reverse=True)
        self.ui.tableWidget_5.setRowCount(len(kee))
        sum = 0
        for i in range(len(kee)):
            self.ui.tableWidget_5.setItem(i, 0, QtWidgets.QTableWidgetItem(str(kee[i])))
            self.ui.tableWidget_5.setItem(i, 1, QtWidgets.QTableWidgetItem(str(k[kee[i]][0])))
            self.ui.tableWidget_5.setItem(i, 2, QtWidgets.QTableWidgetItem(str(k[kee[i]][1])))
            sum += k[kee[i]][1]
        self.ui.label_5.setText('Прибыль: ' + str(sum).replace('.', ','))

    def ras_open(self):
        self.state = 'ras'
        self.ui.gp_in.hide()
        self.s_date = datetime(int(self.ui.l_syear.text()), int(self.ui.l_smonth.text()), int(self.ui.l_sday.text())).date()
        self.n_date = datetime(int(self.ui.l_pyear.text()), int(self.ui.l_pmonth.text()), int(self.ui.l_pday.text())).date()
        self.ui.gp_offer.hide()
        self.ui.groupBox.hide()
        self.ui.gp_services.hide()
        self.ui.gp_stat.hide()
        self.ui.gp_ras.show()
        o = self.ui.tableWidget_4.rowCount()
        for i in range(o):
            self.ui.tableWidget_4.removeRow(0)
        o = load_ras()
        self.ui.tableWidget_4.setRowCount(len(o))
        k = 0
        sum = 0
        for i in range(len(o)):
            date = str(o[i][4]).split('-')
            date = datetime(int(date[0]), int(date[1]), int(date[2])).date()
            if self.s_date <= date <= self.n_date:
                btn = QtWidgets.QPushButton(str(o[i][1]))
                btn.offer_number = o[i][0]
                btn.clicked.connect(self.more_rash)
                self.ui.tableWidget_4.setCellWidget(i - k, 0, btn)
                self.ui.tableWidget_4.setItem(i - k, 1, QtWidgets.QTableWidgetItem(str(o[i][4])))
                self.ui.tableWidget_4.setItem(i - k, 2, QtWidgets.QTableWidgetItem(str(o[i][2])))
                self.ui.tableWidget_4.setItem(i - k, 3, QtWidgets.QTableWidgetItem(str((float(str(o[i][3]).replace(',', '.')) * int(o[i][2]))).replace('.', ',')))
                sum += float(str(o[i][3]).replace(',', '.')) * int(o[i][2])
            else:
                k += 1
        self.ui.l_ras_ito.setText('ИТОГО: ' + str(sum).replace('.', ','))
        self.ui.tableWidget_4.setRowCount(len(o) - k)

    def more_rash(self):
       btn = QtWidgets.QApplication.instance().sender()
       windows.append(More_rash(self, btn.offer_number))
       windows[-1].show()

    def add_rashod(self):
        windows.append(Add_rash(self))
        windows[-1].show()

    def workers_open(self):
        self.state = 'wor'
        self.ui.gp_in.hide()
        self.ui.gp_stat.hide()
        self.ui.gp_ras.hide()
        r = self.ui.tableWidget_2.rowCount()
        for i in range(r):
            self.ui.tableWidget_2.removeRow(i)
        self.ui.tableWidget_2.setRowCount(0)
        self.ui.gp_offer.hide()
        self.ui.gp_services.hide()
        self.ui.groupBox.show()
        workers = load_workers()
        self.ui.tableWidget_2.setRowCount(len(workers))
        for i in range(len(workers)):
            self.ui.tableWidget_2.setItem(i, 0, QtWidgets.QTableWidgetItem(str(workers[i][1].split('ALEEE')[0])))
            self.ui.tableWidget_2.setItem(i, 1, QtWidgets.QTableWidgetItem(str(workers[i][1].split('ALEEE')[1])))
            self.ui.tableWidget_2.setItem(i, 2, QtWidgets.QTableWidgetItem(str(workers[i][1].split('ALEEE')[2])))
            btn = QtWidgets.QPushButton('Подробнее')
            btn.number_of_worker = workers[i][1]
            btn.action_of_btn = 'more_workers'
            btn.clicked.connect(self.more)
            self.ui.tableWidget_2.setCellWidget(i, 3, btn)

    def number_find(self, number=''):
        self.ui.s_page.setValue(1)
        number = self.ui.l_number.text().strip()
        self.offers_open(number_of_client=number)

    def update_page(self):
        offers = len(load_offers())
        if self.on_page * int(self.ui.s_page.text()) - 1 > offers:
            self.ui.s_page.setValue(offers // self.on_page + 1)
        self.offers_open()

    def offers_open(self, number_of_client=''):
        self.state = 'off'
        self.ui.gp_in.hide()
        self.ui.gp_stat.hide()
        self.ui.gp_ras.hide()
        if number_of_client.__class__.__name__ == 'bool':
            number_of_client = ''
        r = self.ui.tableWidget.rowCount()
        for i in range(r):
            self.ui.tableWidget.removeRow(i)
        self.ui.groupBox.hide()
        self.ui.gp_services.hide()
        self.ui.gp_offer.show()
        offers = load_offers()
        offers.reverse()
        offers = list(filter(lambda s: number_of_client in s[8], offers))
        self.on_page = 26
        page = int(self.ui.s_page.text()) - 1
        start = page * self.on_page
        offers = offers[start: start + self.on_page]
        self.ui.tableWidget.setRowCount(len(offers))
        for i in range(len(offers)):
            self.ui.tableWidget.setItem(i, 1, QtWidgets.QTableWidgetItem(str(offers[i][6])))
            if offers[i][7] == 'in':
                col = QtGui.QColor(125, 122, 255)
            elif offers[i][7] == 'co' or offers[i][7] == 'cor':
                col = QtGui.QColor(173, 173, 173)
            elif offers[i][7] == 'wa':
                col = QtGui.QColor(255, 211, 122)
            elif offers[i][7] == 'ne':
                col = QtGui.QColor(122, 255, 127)
            elif offers[i][7] == 'cl':
                col = QtGui.QColor(0, 0, 0)
            elif offers[i][7] == 'nw':
                col = QtGui.QColor(253, 255, 120)
            self.ui.tableWidget.setItem(i, 0, QtWidgets.QTableWidgetItem(str(offers[i][0])))
            self.ui.tableWidget.item(i, 0).setBackground(col)
            self.ui.tableWidget.setItem(i, 2, QtWidgets.QTableWidgetItem(str(offers[i][2])))
            self.ui.tableWidget.setItem(i, 3, QtWidgets.QTableWidgetItem(str(offers[i][3])))
            o = offers[i]
            zzz = o
            o = load_services1(o)
            c = 0
            for io in range(len(o)):
                c += float(o[io][2].replace(',', '.'))
            o = load_sales(zzz)
            for io in range(len(o)):
                c += float(o[io][2].replace(',', '.'))
            self.ui.tableWidget.setItem(i, 4, QtWidgets.QTableWidgetItem(str(c)))
            btn = QtWidgets.QPushButton('Подробнее')
            btn.offer_number = offers[i][0]
            btn.offer = list(filter(lambda s: s[0] == btn.offer_number, load_offers()))
            btn.clicked.connect(self.more_offer)
            self.ui.tableWidget.setCellWidget(i, 5, btn)

    def more_offer(self):
        windows.append(More_offer(self, []))
        windows[-1].show()

    def group_change(self):
        self.services_open()

    def services_open(self):
        self.state = 'serv'
        self.ui.gp_in.hide()
        self.ui.gp_stat.hide()
        r = self.ui.tableWidget_3.rowCount()
        for i in range(r):
            self.ui.tableWidget_3.removeRow(i)
        self.ui.groupBox.hide()
        self.ui.gp_offer.hide()
        self.ui.gp_ras.hide()
        self.ui.gp_services.show()
        if self.ui.service_group.currentText() == 'Все':
            self.ui.pushButton_3.hide()
            services = load_services()
        else:
            self.ui.pushButton_3.show()
            services = load_services(self.ui.service_group.currentText())
        services = list(sorted(services, key=lambda s: s[1]))
        self.ui.tableWidget_3.setRowCount(len(services))
        for i in range(len(services)):
            self.ui.tableWidget_3.setItem(i, 0, QtWidgets.QTableWidgetItem(str(services[i][1])))
            self.ui.tableWidget_3.setItem(i, 1, QtWidgets.QTableWidgetItem(str(services[i][2])))
            bt = QtWidgets.QPushButton('Подробнее')
            bt.number_of_service = services[i][0]
            bt.clicked.connect(self.more_services)
            self.ui.tableWidget_3.setCellWidget(i, 3, bt)
            self.ui.tableWidget_3.setItem(i, 2, QtWidgets.QTableWidgetItem(str(services[i][4])))

    def more_services(self):
        windows.append(More_service(self))
        windows[-1].show()

    def add_services(self):
        windows.append(Add_service(self))
        windows[-1].show()

    def more(self):
        btn = QtWidgets.QApplication.instance().sender()
        if btn.action_of_btn == 'more_workers':
            windows.append(More_workers(btn.number_of_worker))
            windows[-1].show()

    def add_workerss(self):
        windows.append(Add_worker(self))
        windows[-1].show()

    def add_offer(self):
        windows.append(Add_offer(self))
        windows[-1].show()


if (not autho_check()) or today > latest:
    exit()
app = QtWidgets.QApplication([])
application = mywindow()
if passww == []:
    start = First_start_w()
    start.show()
else:
    application.show()

sys.exit(app.exec())
