import docx
from db_working import *


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

doc = docx.Document('offer_list.docx')
# получаем первую таблицу в документе
table = doc.tables[0]
number = list((152, 'g', 'Исузу_эльф', 'Владимир', 'Особые отметки', 'Рекомендации', '2023-11-03', 'nw', '89137344594', None))
if number[9] == None:
    number[9] = ''
cell = table.cell(0, 0)
new_text = cell.text.split('\n')
new_text[0] += str(number[0])
new_text[1] += number[6]
new_text[2] += number[9]
new_text[5] += number[3]
new_text[6] += number[8]
new_text[7] += number[2]
cell.text = '\n'.join(new_text)
table = doc.tables[1]
row = table.rows[-1]
ss = load_services1(number)
for i in range(6):
    remove_row(table, table.rows[1])
o = 0
z = 0
for i in range(len(ss)):
    if ss[i][3] != '0':
        table.add_row()
        k = table.cell(i + 2 - o, 0)
        k.text = str(i + 1 - o)
        k = table.cell(i + 2 - o, 1)
        k.text = ss[i][1]
        k = table.cell(i + 2 - o, 3)
        k.text = ss[i][3]
        k = table.cell(i + 2 - o, 4)
        k.text = str(float(ss[i][2].replace(',', '.')) / int(ss[i][3])).replace('.', ',')
        k = table.cell(i + 2 - o, 5)
        k.text = ss[i][2]
        z += float(ss[i][2].replace(',', '.'))
    else:
        o += 1
k = table.cell(1, 5)
k.text = str(z).replace('.', ',')
row0 = table.rows[1]
row1 = table.rows[-1]
row1._tr.addnext(row0._tr)
p = doc.paragraphs[5]
p.text = f'Расходная накладная к Заказу №{str(number[0])}'

table = doc.tables[2]
ss = load_sales(number)
for i in range(4):
    remove_row(table, table.rows[1])
z = 0
o = 0
for i in range(len(ss)):
    if ss[3] != '0':
        table.add_row()
        k = table.cell(i + 2 - o, 0)
        k.text = str(i + 1 - o)
        k = table.cell(i + 2 - o, 1)
        k.text = ss[i][1]
        k = table.cell(i + 2 - o, 3)
        k.text = ss[i][3]
        k = table.cell(i + 2 - o, 4)
        k.text = str(float(ss[i][2].replace(',', '.')) / int(ss[i][3])).replace('.', ',')
        k = table.cell(i + 2 - o, 5)
        k.text = ss[i][2]
        z += float(ss[i][2].replace(',', '.'))
    else:
        o += 1
k = table.cell(1, 5)
k.text = str(z).replace('.', ',')
row0 = table.rows[1]
row1 = table.rows[-1]
row1._tr.addnext(row0._tr)
table = doc.tables[4]
k = table.cell(0, 0)
k.text = f'Примечания: {number[4]}\nРекомендации: {number[5]}'
doc.save('of.docx')
